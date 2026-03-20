#!/usr/bin/env python3
"""
Resume Adaptation Skill

Adapts a resume to match a specific job description, producing tailored
Word and PDF outputs.

Usage:
    python resume_adapter.py --resume path/to/resume.pdf --job "job description text" --requirements "emphasize leadership"
    python resume_adapter.py --resume resume1.pdf resume2.docx --job https://example.com/job-posting --output-name "google_pm_resume"
"""

import argparse
import os
import re
import sys
from pathlib import Path
from typing import List, Dict, Any, Optional
from dataclasses import dataclass
from datetime import datetime

# Document processing imports
try:
    import docx
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from pypdf import PdfReader
    PDF_AVAILABLE = True
except ImportError:
    try:
        from PyPDF2 import PdfReader
        PDF_AVAILABLE = True
    except ImportError:
        PDF_AVAILABLE = False

try:
    import markdown
    MARKDOWN_AVAILABLE = True
except ImportError:
    MARKDOWN_AVAILABLE = False

try:
    from docx2pdf import convert
    DOCX2PDF_AVAILABLE = True
except ImportError:
    DOCX2PDF_AVAILABLE = False


@dataclass
class JobRequirements:
    """Structured representation of job requirements."""
    title: str = ""
    company: str = ""
    required_skills: List[str] = None
    preferred_skills: List[str] = None
    responsibilities: List[str] = None
    qualifications: List[str] = None
    experience_years: Optional[int] = None
    education: List[str] = None
    keywords: List[str] = None

    def __post_init__(self):
        if self.required_skills is None:
            self.required_skills = []
        if self.preferred_skills is None:
            self.preferred_skills = []
        if self.responsibilities is None:
            self.responsibilities = []
        if self.qualifications is None:
            self.qualifications = []
        if self.education is None:
            self.education = []
        if self.keywords is None:
            self.keywords = []


@dataclass
class ResumeSection:
    """A section of a resume."""
    title: str
    content: str
    order: int = 0


@dataclass
class ParsedResume:
    """Structured representation of a parsed resume."""
    name: str = ""
    contact_info: Dict[str, str] = None
    summary: str = ""
    sections: List[ResumeSection] = None
    skills: List[str] = None
    experience: List[Dict[str, Any]] = None
    education: List[Dict[str, Any]] = None
    raw_text: str = ""

    def __post_init__(self):
        if self.contact_info is None:
            self.contact_info = {}
        if self.sections is None:
            self.sections = []
        if self.skills is None:
            self.skills = []
        if self.experience is None:
            self.experience = []
        if self.education is None:
            self.education = []


class DocumentParser:
    """Parser for various document formats."""

    @staticmethod
    def parse_pdf(file_path: str) -> str:
        """Extract text from a PDF file."""
        if not PDF_AVAILABLE:
            raise ImportError("PDF parsing requires pypdf or PyPDF2. Install with: pip install pypdf")

        reader = PdfReader(file_path)
        text = ""
        for page in reader.pages:
            text += page.extract_text() + "\n"
        return text

    @staticmethod
    def parse_docx(file_path: str) -> str:
        """Extract text from a DOCX file."""
        if not DOCX_AVAILABLE:
            raise ImportError("DOCX parsing requires python-docx. Install with: pip install python-docx")

        doc = Document(file_path)
        text = ""
        for para in doc.paragraphs:
            text += para.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + "\n"
        return text

    @staticmethod
    def parse_txt(file_path: str) -> str:
        """Read text from a plain text file."""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()

    @staticmethod
    def parse_md(file_path: str) -> str:
        """Read and optionally convert markdown to text."""
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        if MARKDOWN_AVAILABLE:
            # Convert to HTML, then we can extract text if needed
            return markdown.markdown(content)
        return content

    @classmethod
    def parse(cls, file_path: str) -> str:
        """Parse any supported file type."""
        ext = Path(file_path).suffix.lower()
        parsers = {
            '.pdf': cls.parse_pdf,
            '.docx': cls.parse_docx,
            '.txt': cls.parse_txt,
            '.md': cls.parse_md,
        }

        if ext not in parsers:
            raise ValueError(f"Unsupported file format: {ext}")

        return parsers[ext](file_path)


class ResumeAnalyzer:
    """Analyzes and structures resume content."""

    SECTION_HEADERS = [
        'summary', 'professional summary', 'profile', 'objective',
        'experience', 'work experience', 'professional experience', 'employment',
        'education', 'academic background',
        'skills', 'technical skills', 'core competencies', 'expertise',
        'projects', 'certifications', 'awards', 'publications',
        'languages', 'interests', 'references'
    ]

    def analyze(self, raw_text: str) -> ParsedResume:
        """Convert raw text into a structured resume object."""
        resume = ParsedResume(raw_text=raw_text)

        # Split into lines and identify sections
        lines = raw_text.split('\n')

        # Try to extract name (usually first non-empty line)
        for line in lines:
            line = line.strip()
            if line and not line.lower() in self.SECTION_HEADERS:
                resume.name = line
                break

        # Extract contact info patterns
        resume.contact_info = self._extract_contact_info(raw_text)

        # Identify and extract sections
        resume.sections = self._extract_sections(raw_text)

        # Extract skills
        resume.skills = self._extract_skills(raw_text)

        # Extract experience
        resume.experience = self._extract_experience(raw_text)

        # Extract education
        resume.education = self._extract_education(raw_text)

        # Find summary section
        for section in resume.sections:
            if section.title.lower() in ['summary', 'professional summary', 'profile']:
                resume.summary = section.content
                break

        return resume

    def _extract_contact_info(self, text: str) -> Dict[str, str]:
        """Extract contact information from resume text."""
        contact = {}

        # Email pattern
        email_match = re.search(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', text)
        if email_match:
            contact['email'] = email_match.group(0)

        # Phone pattern
        phone_match = re.search(r'(\+?1[-.\s]?)?\(?[0-9]{3}\)?[-.\s]?[0-9]{3}[-.\s]?[0-9]{4}', text)
        if phone_match:
            contact['phone'] = phone_match.group(0)

        # LinkedIn
        linkedin_match = re.search(r'linkedin\.com/in/[A-Za-z0-9_-]+', text)
        if linkedin_match:
            contact['linkedin'] = linkedin_match.group(0)

        # Website/Portfolio
        website_match = re.search(r'https?://[^\s]+|www\.[^\s]+', text)
        if website_match:
            contact['website'] = website_match.group(0)

        return contact

    def _extract_sections(self, text: str) -> List[ResumeSection]:
        """Extract sections from resume text."""
        sections = []
        lines = text.split('\n')

        current_section = None
        current_content = []
        section_order = 0

        for line in lines:
            stripped = line.strip()
            if not stripped:
                continue

            # Check if this line is a section header
            is_header = False
            for header in self.SECTION_HEADERS:
                if stripped.lower() == header or stripped.lower().rstrip(':') == header:
                    if current_section:
                        sections.append(ResumeSection(
                            title=current_section,
                            content='\n'.join(current_content),
                            order=section_order
                        ))
                    current_section = stripped
                    current_content = []
                    section_order += 1
                    is_header = True
                    break

            if not is_header and current_section:
                current_content.append(line)

        # Don't forget the last section
        if current_section and current_content:
            sections.append(ResumeSection(
                title=current_section,
                content='\n'.join(current_content),
                order=section_order
            ))

        return sections

    def _extract_skills(self, text: str) -> List[str]:
        """Extract skills from resume text."""
        skills = []

        # Look for skills section
        skills_section_match = re.search(
            r'(?:skills|technical skills|core competencies|expertise)[\s:]*\n(.*?)(?:\n\n|\n[A-Z]|$)',
            text, re.IGNORECASE | re.DOTALL
        )

        if skills_section_match:
            skills_text = skills_section_match.group(1)
            # Split by common delimiters
            skills = [s.strip() for s in re.split(r'[,•|\n]+', skills_text) if s.strip()]

        return skills

    def _extract_experience(self, text: str) -> List[Dict[str, Any]]:
        """Extract work experience entries."""
        experiences = []

        # This is a simplified extraction - real implementation would be more sophisticated
        experience_section_match = re.search(
            r'(?:experience|work experience|professional experience)[\s:]*\n(.*?)(?:\n(?:education|skills|projects|certifications)|$)',
            text, re.IGNORECASE | re.DOTALL
        )

        if experience_section_match:
            exp_text = experience_section_match.group(1)
            # Split by double newlines to get individual jobs
            jobs = re.split(r'\n\n+', exp_text)
            for job in jobs:
                if job.strip():
                    experiences.append({
                        'raw': job.strip(),
                        'title': '',
                        'company': '',
                        'dates': '',
                        'description': job.strip()
                    })

        return experiences

    def _extract_education(self, text: str) -> List[Dict[str, Any]]:
        """Extract education entries."""
        education = []

        education_section_match = re.search(
            r'(?:education|academic background)[\s:]*\n(.*?)(?:\n\n[A-Z]|$)',
            text, re.IGNORECASE | re.DOTALL
        )

        if education_section_match:
            edu_text = education_section_match.group(1)
            entries = re.split(r'\n\n+', edu_text)
            for entry in entries:
                if entry.strip():
                    education.append({
                        'raw': entry.strip(),
                        'institution': '',
                        'degree': '',
                        'dates': '',
                        'details': entry.strip()
                    })

        return education


class JobDescriptionAnalyzer:
    """Analyzes job descriptions to extract requirements."""

    def analyze(self, text: str) -> JobRequirements:
        """Extract structured requirements from job description."""
        requirements = JobRequirements()

        # Extract job title (usually first line or in first paragraph)
        lines = [l.strip() for l in text.split('\n') if l.strip()]
        if lines:
            requirements.title = lines[0][:100]  # First line, truncated

        # Extract skills using common patterns
        requirements.required_skills = self._extract_skills(text, 'required')
        requirements.preferred_skills = self._extract_skills(text, 'preferred')

        # Combine all skills for keyword matching
        all_skills = requirements.required_skills + requirements.preferred_skills

        # Extract responsibilities
        requirements.responsibilities = self._extract_responsibilities(text)

        # Extract qualifications
        requirements.qualifications = self._extract_qualifications(text)

        # Extract experience years
        requirements.experience_years = self._extract_experience_years(text)

        # Build keyword list
        requirements.keywords = self._extract_keywords(text, all_skills)

        return requirements

    def _extract_skills(self, text: str, skill_type: str = 'all') -> List[str]:
        """Extract skills from job description."""
        skills = []

        # Common technical skills to look for
        tech_skills = [
            'python', 'javascript', 'java', 'c++', 'c#', 'go', 'rust', 'ruby',
            'react', 'angular', 'vue', 'node.js', 'django', 'flask', 'spring',
            'aws', 'azure', 'gcp', 'docker', 'kubernetes', 'terraform',
            'sql', 'postgresql', 'mysql', 'mongodb', 'redis', 'elasticsearch',
            'machine learning', 'ai', 'data science', 'analytics',
            'agile', 'scrum', 'kanban', 'jira', 'confluence',
            'leadership', 'management', 'communication', 'collaboration'
        ]

        text_lower = text.lower()
        for skill in tech_skills:
            if skill in text_lower:
                skills.append(skill.title() if skill != 'ai' else 'AI')

        # Look for explicit skill lists
        if skill_type == 'required':
            pattern = r'(?:required skills?|requirements?|must have|required)[\s:]*\n?(.*?)(?:preferred|nice to have|qualifications|$)'
        elif skill_type == 'preferred':
            pattern = r'(?:preferred skills?|nice to have|preferred qualifications?)[\s:]*\n?(.*?)(?:requirements?|$)'
        else:
            pattern = r'(?:skills? required|technical skills?)[\s:]*\n?(.*?)(?:\n\n|$)'

        match = re.search(pattern, text, re.IGNORECASE | re.DOTALL)
        if match:
            skills_text = match.group(1)
            found_skills = [s.strip() for s in re.split(r'[,•|\n]+', skills_text) if s.strip() and len(s.strip()) < 50]
            skills.extend(found_skills)

        return list(set(skills))  # Remove duplicates

    def _extract_responsibilities(self, text: str) -> List[str]:
        """Extract job responsibilities."""
        responsibilities = []

        pattern = r'(?:responsibilities?|what you.ll do|key duties|role overview)[\s:]*\n?(.*?)(?:requirements?|qualifications?|what you bring|$)'
        match = re.search(pattern, text, re.IGNORECASE | re.DOTALL)

        if match:
            resp_text = match.group(1)
            # Split by bullet points or numbers
            items = re.split(r'[•\-\*]|\d+\.', resp_text)
            responsibilities = [item.strip() for item in items if len(item.strip()) > 10 and len(item.strip()) < 200]

        return responsibilities[:10]  # Limit to top 10

    def _extract_qualifications(self, text: str) -> List[str]:
        """Extract qualifications."""
        qualifications = []

        pattern = r'(?:qualifications?|requirements?|what you bring|what we.looking for)[\s:]*\n?(.*?)(?:benefits?|perks?|about us|$)'
        match = re.search(pattern, text, re.IGNORECASE | re.DOTALL)

        if match:
            qual_text = match.group(1)
            items = re.split(r'[•\-\*]|\d+\.', qual_text)
            qualifications = [item.strip() for item in items if len(item.strip()) > 10 and len(item.strip()) < 200]

        return qualifications[:10]

    def _extract_experience_years(self, text: str) -> Optional[int]:
        """Extract required years of experience."""
        patterns = [
            r'(\d+)\+?\s*years?\s*of\s*experience',
            r'(\d+)\+?\s*years?\s*experience',
            r'(\d+)\+?\s*years?\s*(?:in|of)',
            r'minimum\s*of\s*(\d+)\s*years?'
        ]

        text_lower = text.lower()
        for pattern in patterns:
            match = re.search(pattern, text_lower)
            if match:
                return int(match.group(1))

        return None

    def _extract_keywords(self, text: str, skills: List[str]) -> List[str]:
        """Extract important keywords for matching."""
        keywords = set(skills)

        # Add action verbs and important terms
        important_terms = [
            'leadership', 'management', 'strategy', 'architecture', 'design',
            'development', 'implementation', 'optimization', 'analysis',
            'collaboration', 'stakeholder', 'cross-functional', 'mentoring'
        ]

        text_lower = text.lower()
        for term in important_terms:
            if term in text_lower:
                keywords.add(term)

        return list(keywords)


class ResumeTailor:
    """Tailors a resume to match job requirements."""

    def __init__(self, resume: ParsedResume, job: JobRequirements, user_requirements: str = ""):
        self.resume = resume
        self.job = job
        self.user_requirements = user_requirements

    def tailor(self) -> ParsedResume:
        """Create a tailored version of the resume."""
        tailored = ParsedResume()

        # Copy basic info
        tailored.name = self.resume.name
        tailored.contact_info = self.resume.contact_info.copy()

        # Tailor summary
        tailored.summary = self._tailor_summary()

        # Tailor skills section
        tailored.skills = self._prioritize_skills()

        # Tailor experience
        tailored.experience = self._tailor_experience()

        # Tailor education
        tailored.education = self.resume.education

        # Build tailored sections
        tailored.sections = self._build_sections(tailored)

        return tailored

    def _tailor_summary(self) -> str:
        """Create a tailored professional summary."""
        if not self.resume.summary:
            # Generate a summary if none exists
            return self._generate_summary()

        # Enhance existing summary with job keywords
        summary = self.resume.summary

        # Add relevant keywords that appear in job but not in summary
        summary_lower = summary.lower()
        keywords_to_add = []

        for keyword in self.job.keywords[:5]:
            if keyword.lower() not in summary_lower and len(keywords_to_add) < 3:
                keywords_to_add.append(keyword)

        if keywords_to_add:
            summary += f" Experienced in {', '.join(keywords_to_add)}."

        return summary

    def _generate_summary(self) -> str:
        """Generate a professional summary based on resume content."""
        name_parts = self.resume.name.split()
        first_name = name_parts[0] if name_parts else "Professional"

        # Get top skills that match job
        matching_skills = [s for s in self.resume.skills if any(
            js.lower() in s.lower() or s.lower() in js.lower()
            for js in self.job.required_skills
        )]

        top_skills = matching_skills[:5] if matching_skills else self.resume.skills[:5]

        summary = f"{first_name} is a results-driven professional"

        if self.job.title:
            summary += f" with expertise in {self.job.title.lower()}"

        if top_skills:
            summary += f". Skilled in {', '.join(top_skills)}"

        summary += " with a proven track record of delivering impactful results."

        return summary

    def _prioritize_skills(self) -> List[str]:
        """Prioritize skills based on job requirements."""
        # Start with matching skills
        prioritized = []
        remaining = []

        for skill in self.resume.skills:
            is_match = any(
                job_skill.lower() in skill.lower() or skill.lower() in job_skill.lower()
                for job_skill in self.job.required_skills + self.job.preferred_skills
            )
            if is_match:
                prioritized.append(skill)
            else:
                remaining.append(skill)

        # Combine prioritized first, then remaining
        return prioritized + remaining

    def _tailor_experience(self) -> List[Dict[str, Any]]:
        """Reorder and enhance experience entries based on relevance."""
        scored_experiences = []

        for exp in self.resume.experience:
            score = self._score_experience_relevance(exp)
            scored_experiences.append((score, exp))

        # Sort by relevance score (descending)
        scored_experiences.sort(key=lambda x: x[0], reverse=True)

        return [exp for _, exp in scored_experiences]

    def _score_experience_relevance(self, exp: Dict[str, Any]) -> int:
        """Score how relevant an experience entry is to the job."""
        score = 0
        exp_text = exp.get('raw', '').lower()

        # Check for keyword matches
        for keyword in self.job.keywords:
            if keyword.lower() in exp_text:
                score += 2

        # Check for skill matches
        for skill in self.job.required_skills:
            if skill.lower() in exp_text:
                score += 3

        for skill in self.job.preferred_skills:
            if skill.lower() in exp_text:
                score += 1

        return score

    def _build_sections(self, tailored: ParsedResume) -> List[ResumeSection]:
        """Build the final sections for the tailored resume."""
        sections = []
        order = 0

        # Summary section
        if tailored.summary:
            sections.append(ResumeSection(
                title="Professional Summary",
                content=tailored.summary,
                order=order
            ))
            order += 1

        # Skills section
        if tailored.skills:
            skills_text = "• " + "\n• ".join(tailored.skills)
            sections.append(ResumeSection(
                title="Skills",
                content=skills_text,
                order=order
            ))
            order += 1

        # Experience section
        if tailored.experience:
            exp_content = "\n\n".join([exp.get('description', '') for exp in tailored.experience])
            sections.append(ResumeSection(
                title="Professional Experience",
                content=exp_content,
                order=order
            ))
            order += 1

        # Education section
        if tailored.education:
            edu_content = "\n\n".join([edu.get('details', '') for edu in tailored.education])
            sections.append(ResumeSection(
                title="Education",
                content=edu_content,
                order=order
            ))
            order += 1

        return sections


class ATSScorer:
    """Scores a resume against ATS (Applicant Tracking System) requirements."""

    def __init__(self, resume: ParsedResume, job: JobRequirements):
        self.resume = resume
        self.job = job
        self.resume_text_lower = resume.raw_text.lower()

    def score(self) -> Dict[str, Any]:
        """Calculate comprehensive ATS compatibility score."""
        scores = {
            'overall': 0,
            'sections': {},
            'suggestions': [],
            'keyword_match': {},
            'format_score': 0,
            'content_score': 0
        }

        # Score individual components
        scores['sections']['skills_match'] = self._score_skills_match()
        scores['sections']['experience_match'] = self._score_experience_match()
        scores['sections']['education_match'] = self._score_education_match()
        scores['sections']['keyword_density'] = self._score_keyword_density()
        scores['format_score'] = self._score_format()
        scores['content_score'] = self._score_content()

        # Calculate overall score (weighted average)
        weights = {
            'skills_match': 0.25,
            'experience_match': 0.25,
            'education_match': 0.15,
            'keyword_density': 0.20,
        }

        overall = sum(
            scores['sections'][section] * weight
            for section, weight in weights.items()
        )
        # Add format and content scores directly
        overall += scores['format_score'] * 0.10
        overall += scores['content_score'] * 0.05

        scores['overall'] = round(overall, 1)

        # Generate improvement suggestions
        scores['suggestions'] = self._generate_suggestions(scores)

        # Detailed keyword matching
        scores['keyword_match'] = self._analyze_keyword_match()

        return scores

    def _score_skills_match(self) -> float:
        """Score how well resume skills match job requirements."""
        if not self.job.required_skills and not self.job.preferred_skills:
            return 50.0  # Neutral if no skills specified

        required_matches = 0
        preferred_matches = 0

        # Check required skills (worth more)
        for skill in self.job.required_skills:
            if self._skill_in_resume(skill):
                required_matches += 1

        # Check preferred skills
        for skill in self.job.preferred_skills:
            if self._skill_in_resume(skill):
                preferred_matches += 1

        # Calculate scores
        required_score = (required_matches / len(self.job.required_skills) * 100) if self.job.required_skills else 0
        preferred_score = (preferred_matches / len(self.job.preferred_skills) * 100) if self.job.preferred_skills else 0

        # Weight required more heavily
        if self.job.required_skills and self.job.preferred_skills:
            return round((required_score * 0.7) + (preferred_score * 0.3), 1)
        elif self.job.required_skills:
            return round(required_score, 1)
        else:
            return round(preferred_score, 1)

    def _skill_in_resume(self, skill: str) -> bool:
        """Check if a skill appears in the resume."""
        skill_lower = skill.lower()

        # Check in skills list
        for resume_skill in self.resume.skills:
            if skill_lower in resume_skill.lower() or resume_skill.lower() in skill_lower:
                return True

        # Check in full resume text
        if skill_lower in self.resume_text_lower:
            return True

        return False

    def _score_experience_match(self) -> float:
        """Score how well experience matches job requirements."""
        if not self.job.responsibilities:
            return 50.0

        matches = 0
        for resp in self.job.responsibilities:
            resp_keywords = self._extract_keywords_from_text(resp.lower())
            if any(kw in self.resume_text_lower for kw in resp_keywords):
                matches += 1

        return round((matches / len(self.job.responsibilities)) * 100, 1) if self.job.responsibilities else 50.0

    def _score_education_match(self) -> float:
        """Score education requirements match."""
        # Check for degree mentions
        degree_patterns = [
            r'bachelor', r'master', r'phd', r'doctorate', r'bs', r'ms', r'ba', r'ma',
            r'computer science', r'engineering', r'business', r'mba'
        ]

        education_text = ' '.join([edu.get('details', '').lower() for edu in self.resume.education])

        matches = sum(1 for pattern in degree_patterns if pattern in education_text)
        return min(100, round((matches / 4) * 100, 1))  # Cap at 100

    def _score_keyword_density(self) -> float:
        """Score keyword density and distribution."""
        if not self.job.keywords:
            return 50.0

        found_keywords = []
        for keyword in self.job.keywords:
            count = self.resume_text_lower.count(keyword.lower())
            if count > 0:
                found_keywords.append((keyword, count))

        # Calculate density score
        coverage = len(found_keywords) / len(self.job.keywords) if self.job.keywords else 0

        # Check distribution (keywords should appear in different sections)
        sections_with_keywords = set()
        for section in self.resume.sections:
            section_text = section.content.lower()
            for keyword, _ in found_keywords:
                if keyword.lower() in section_text:
                    sections_with_keywords.add(section.title)

        distribution_score = len(sections_with_keywords) / len(self.resume.sections) if self.resume.sections else 0

        return round((coverage * 0.7 + distribution_score * 0.3) * 100, 1)

    def _score_format(self) -> float:
        """Score resume format for ATS compatibility."""
        score = 100

        # Check for common ATS-unfriendly elements
        issues = []

        # No contact info
        if not self.resume.contact_info.get('email') or not self.resume.contact_info.get('phone'):
            score -= 20
            issues.append("Missing contact information")

        # No clear section headers
        if len(self.resume.sections) < 3:
            score -= 15
            issues.append("Limited section structure")

        # Very long resume (ATS may truncate)
        word_count = len(self.resume_text_lower.split())
        if word_count > 1000:
            score -= 10
            issues.append("Resume may be too long for some ATS systems")

        # Check for special characters that might confuse ATS
        special_chars = re.findall(r'[^\w\s\-\.\@\(\)\,\|/]', self.resume.raw_text)
        if len(special_chars) > 20:
            score -= 10
            issues.append("Many special characters may confuse ATS")

        return max(0, round(score, 1))

    def _score_content(self) -> float:
        """Score content quality for ATS."""
        score = 100

        # Check for action verbs
        action_verbs = [
            'achieved', 'improved', 'trained', 'managed', 'created', 'developed',
            'led', 'implemented', 'designed', 'launched', 'increased', 'decreased',
            'negotiated', 'supervised', 'generated', 'delivered', 'established'
        ]

        verb_count = sum(1 for verb in action_verbs if verb in self.resume_text_lower)
        if verb_count < 3:
            score -= 15

        # Check for measurable results (numbers)
        numbers = re.findall(r'\d+%|\$\d+|\d+\s*(million|billion|k|thousand)', self.resume_text_lower)
        if len(numbers) < 2:
            score -= 10

        # Check for passive voice indicators
        passive_indicators = ['was responsible', 'was tasked', 'was involved', 'assisted with']
        passive_count = sum(1 for indicator in passive_indicators if indicator in self.resume_text_lower)
        if passive_count > 3:
            score -= 10

        return max(0, round(score, 1))

    def _generate_suggestions(self, scores: Dict) -> List[str]:
        """Generate improvement suggestions based on scores."""
        suggestions = []

        if scores['sections']['skills_match'] < 70:
            suggestions.append(
                f"Add more relevant skills. Missing: {', '.join(self._get_missing_skills()[:5])}"
            )

        if scores['sections']['keyword_density'] < 60:
            suggestions.append(
                "Increase keyword density by incorporating more job-specific terms naturally throughout your resume"
            )

        if scores['format_score'] < 80:
            suggestions.append(
                "Simplify formatting: use standard section headers, avoid tables and graphics, ensure contact info is complete"
            )

        if scores['content_score'] < 70:
            suggestions.append(
                "Use more action verbs and quantify achievements (e.g., 'Increased sales by 25%')"
            )

        if not suggestions:
            suggestions.append("Great job! Your resume is well-optimized for ATS systems.")

        return suggestions

    def _get_missing_skills(self) -> List[str]:
        """Get list of job skills not found in resume."""
        missing = []
        all_job_skills = self.job.required_skills + self.job.preferred_skills

        for skill in all_job_skills:
            if not self._skill_in_resume(skill):
                missing.append(skill)

        return missing

    def _analyze_keyword_match(self) -> Dict[str, List[str]]:
        """Analyze which keywords match and which are missing."""
        return {
            'found': [kw for kw in self.job.keywords if kw.lower() in self.resume_text_lower],
            'missing': [kw for kw in self.job.keywords if kw.lower() not in self.resume_text_lower],
            'required_found': [s for s in self.job.required_skills if self._skill_in_resume(s)],
            'required_missing': [s for s in self.job.required_skills if not self._skill_in_resume(s)],
        }

    def _extract_keywords_from_text(self, text: str) -> List[str]:
        """Extract keywords from a text string."""
        # Simple keyword extraction - split and filter
        words = re.findall(r'\b[a-z]{4,}\b', text)
        return list(set(words))


class CoverLetterGenerator:
    """Generates a tailored cover letter based on resume and job description."""

    def __init__(self, resume: ParsedResume, job: JobRequirements, user_requirements: str = ""):
        self.resume = resume
        self.job = job
        self.user_requirements = user_requirements

    def generate(self) -> str:
        """Generate a complete cover letter."""
        sections = []

        # Header with contact info
        sections.append(self._generate_header())

        # Salutation
        sections.append(self._generate_salutation())

        # Opening paragraph
        sections.append(self._generate_opening())

        # Body paragraphs
        sections.append(self._generate_body())

        # Closing paragraph
        sections.append(self._generate_closing())

        # Sign-off
        sections.append(self._generate_signoff())

        return '\n\n'.join(sections)

    def _generate_header(self) -> str:
        """Generate cover letter header with contact information."""
        name = self.resume.name
        contact_parts = []

        if self.resume.contact_info.get('email'):
            contact_parts.append(self.resume.contact_info['email'])
        if self.resume.contact_info.get('phone'):
            contact_parts.append(self.resume.contact_info['phone'])
        if self.resume.contact_info.get('linkedin'):
            contact_parts.append(self.resume.contact_info['linkedin'])
        if self.resume.contact_info.get('website'):
            contact_parts.append(self.resume.contact_info['website'])

        header = name + '\n'
        if contact_parts:
            header += ' | '.join(contact_parts)

        # Add date
        from datetime import datetime
        header += f"\n\n{datetime.now().strftime('%B %d, %Y')}"

        return header

    def _generate_salutation(self) -> str:
        """Generate appropriate salutation."""
        # Default to generic if no specific hiring manager known
        return "Dear Hiring Manager,"

    def _generate_opening(self) -> str:
        """Generate opening paragraph."""
        job_title = self.job.title or "the position"
        company = self.job.company or "your company"

        # Get top matching skills
        matching_skills = self._get_top_matching_skills(3)

        opening = f"I am writing to express my strong interest in {job_title} at {company}. "
        opening += f"With my background in {', '.join(matching_skills)}, "
        opening += "I am excited about the opportunity to contribute to your team."

        # Incorporate user requirements if relevant
        if self.user_requirements:
            opening += f" {self.user_requirements}"

        return opening

    def _generate_body(self) -> str:
        """Generate body paragraphs highlighting relevant experience."""
        paragraphs = []

        # Paragraph 1: Key qualifications and matching experience
        para1 = "My experience has prepared me well for this role. "

        # Highlight most relevant experience
        if self.resume.experience:
            top_exp = self.resume.experience[0]
            exp_desc = top_exp.get('description', '')
            # Extract first sentence or truncate
            first_sentence = exp_desc.split('.')[0] if '.' in exp_desc else exp_desc[:150]
            para1 += f"{first_sentence}. "

        # Mention specific job requirements
        if self.job.responsibilities:
            resp_matches = []
            for resp in self.job.responsibilities[:2]:
                resp_keywords = self._extract_keywords(resp)
                if any(kw.lower() in self.resume.raw_text.lower() for kw in resp_keywords):
                    resp_matches.append(resp)

            if resp_matches:
                para1 += f"I have demonstrated success in similar responsibilities, including {resp_matches[0].lower()}."

        paragraphs.append(para1)

        # Paragraph 2: Skills alignment
        para2 = "My skill set aligns closely with your requirements. "

        if self.job.required_skills:
            matched_skills = [s for s in self.job.required_skills if self._skill_in_resume(s)]
            if matched_skills:
                para2 += f"I bring expertise in {', '.join(matched_skills[:4])}. "

        # Add soft skills if mentioned
        soft_skills = ['leadership', 'communication', 'collaboration', 'problem-solving']
        found_soft = [s for s in soft_skills if s in self.resume.raw_text.lower()]
        if found_soft:
            para2 += f"Additionally, my {found_soft[0]} skills have enabled me to work effectively in cross-functional teams and deliver results."

        paragraphs.append(para2)

        # Paragraph 3: Value proposition
        para3 = "I am particularly drawn to this opportunity because of "
        para3 += "the potential to make a meaningful impact. "
        para3 += "I am confident that my combination of technical expertise and "
        para3 += "practical experience makes me a strong fit for your team."

        paragraphs.append(para3)

        return '\n\n'.join(paragraphs)

    def _generate_closing(self) -> str:
        """Generate closing paragraph."""
        closing = "Thank you for considering my application. I would welcome the opportunity "
        closing += "to discuss how my background and skills can contribute to your team's success. "
        closing += "I look forward to the possibility of speaking with you soon."
        return closing

    def _generate_signoff(self) -> str:
        """Generate sign-off."""
        name_parts = self.resume.name.split()
        first_name = name_parts[0] if name_parts else ""

        return f"Sincerely,\n\n{first_name}"

    def _get_top_matching_skills(self, count: int = 3) -> List[str]:
        """Get top skills that match the job."""
        matches = []
        all_job_skills = self.job.required_skills + self.job.preferred_skills

        for skill in all_job_skills:
            if self._skill_in_resume(skill) and len(matches) < count:
                matches.append(skill)

        # Fill with general skills if needed
        if len(matches) < count and self.resume.skills:
            matches.extend(self.resume.skills[:count - len(matches)])

        return matches if matches else ["relevant areas"]

    def _skill_in_resume(self, skill: str) -> bool:
        """Check if a skill appears in the resume."""
        skill_lower = skill.lower()
        resume_lower = self.resume.raw_text.lower()

        for resume_skill in self.resume.skills:
            if skill_lower in resume_skill.lower() or resume_skill.lower() in skill_lower:
                return True

        return skill_lower in resume_lower

    def _extract_keywords(self, text: str) -> List[str]:
        """Extract meaningful keywords from text."""
        words = re.findall(r'\b[a-zA-Z]{4,}\b', text)
        return [w.lower() for w in words]


class ResumeDocumentBuilder:
    """Builds DOCX and PDF documents from a tailored resume."""

    def __init__(self, output_dir: str = "./output"):
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)

    def build_docx(self, resume: ParsedResume, output_name: str) -> str:
        """Create a DOCX file from the tailored resume."""
        if not DOCX_AVAILABLE:
            raise ImportError("DOCX creation requires python-docx. Install with: pip install python-docx")

        doc = Document()

        # Set up styles
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)

        # Add name as title
        title = doc.add_heading(resume.name, level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add contact info
        contact_line = " | ".join([
            resume.contact_info.get('email', ''),
            resume.contact_info.get('phone', ''),
            resume.contact_info.get('linkedin', ''),
            resume.contact_info.get('website', '')
        ])
        contact_line = ' | '.join(filter(None, contact_line.split(' | ')))

        if contact_line:
            contact_para = doc.add_paragraph(contact_line)
            contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()  # Spacing

        # Add sections
        for section in sorted(resume.sections, key=lambda s: s.order):
            heading = doc.add_heading(section.title, level=1)

            # Add content
            if section.title.lower() == 'skills':
                # Skills as a paragraph with bullets
                skills_para = doc.add_paragraph()
                skills = section.content.replace('• ', '').split('\n')
                skills_para.add_run(', '.join([s.strip() for s in skills if s.strip()]))
            else:
                # Other sections as paragraphs
                for line in section.content.split('\n'):
                    if line.strip():
                        if line.strip().startswith('•'):
                            doc.add_paragraph(line.strip(), style='List Bullet')
                        else:
                            doc.add_paragraph(line.strip())

            doc.add_paragraph()  # Spacing between sections

        # Save document
        output_path = self.output_dir / f"{output_name}.docx"
        doc.save(output_path)
        return str(output_path)

    def build_pdf(self, docx_path: str, output_name: str) -> str:
        """Convert DOCX to PDF using reportlab."""
        output_path = self.output_dir / f"{output_name}.pdf"

        if DOCX2PDF_AVAILABLE:
            # Use docx2pdf library (Windows/Mac with Word installed)
            convert(docx_path, str(output_path))
        else:
            # Convert DOCX to PDF using reportlab
            self._convert_docx_to_pdf_with_reportlab(docx_path, str(output_path))

        return str(output_path)

    def _convert_docx_to_pdf_with_reportlab(self, docx_path: str, pdf_path: str):
        """Convert DOCX to PDF using reportlab by extracting and recreating content."""
        from reportlab.lib.pagesizes import letter
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.enums import TA_CENTER, TA_LEFT
        from reportlab.lib.units import inch

        # Read the DOCX file
        doc = Document(docx_path)

        # Create PDF document
        pdf_doc = SimpleDocTemplate(
            pdf_path,
            pagesize=letter,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=18
        )

        # Container for the 'Flowable' objects
        elements = []

        # Define styles
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=18,
            textColor='#000000',
            spaceAfter=12,
            alignment=TA_CENTER,
            fontName='Helvetica-Bold'
        )
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=12,
            textColor='#000000',
            spaceAfter=6,
            spaceBefore=12,
            fontName='Helvetica-Bold'
        )
        normal_style = ParagraphStyle(
            'CustomNormal',
            parent=styles['Normal'],
            fontSize=10,
            leading=14,
            alignment=TA_LEFT
        )
        center_style = ParagraphStyle(
            'CustomCenter',
            parent=styles['Normal'],
            fontSize=10,
            leading=14,
            alignment=TA_CENTER
        )

        # Process paragraphs
        is_first = True
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                elements.append(Spacer(1, 6))
                continue

            # Check if this is the title (first paragraph, usually name)
            if is_first:
                elements.append(Paragraph(text, title_style))
                is_first = False
            # Check if this looks like a section heading
            elif text.isupper() or (len(text) < 50 and text.endswith(':')):
                elements.append(Paragraph(text, heading_style))
            # Center-aligned contact info
            elif '@' in text or 'linkedin.com' in text or any(c.isdigit() for c in text.replace(' ', '')) and len(text) < 100:
                elements.append(Paragraph(text, center_style))
            else:
                elements.append(Paragraph(text, normal_style))

        # Process tables (if any)
        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                if row_text:
                    elements.append(Paragraph(row_text, normal_style))
            elements.append(Spacer(1, 6))

        # Build PDF
        pdf_doc.build(elements)

    def build_cover_letter_docx(self, cover_letter_text: str, output_name: str) -> str:
        """Create a DOCX file for the cover letter."""
        if not DOCX_AVAILABLE:
            raise ImportError("DOCX creation requires python-docx. Install with: pip install python-docx")

        doc = Document()

        # Set up styles
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)

        # Parse cover letter sections
        sections = cover_letter_text.split('\n\n')

        for section in sections:
            if not section.strip():
                continue

            # Check if this is the header (contains contact info)
            if '|' in section or '@' in section[:100]:
                # Header section
                for line in section.split('\n'):
                    if line.strip():
                        para = doc.add_paragraph(line.strip())
                        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            else:
                # Regular paragraph
                para = doc.add_paragraph(section.strip())
                para_format = para.paragraph_format
                para_format.space_after = Pt(12)

        # Save document
        output_path = self.output_dir / f"{output_name}_cover_letter.docx"
        doc.save(output_path)
        return str(output_path)

    def build_cover_letter_pdf(self, docx_path: str, output_name: str) -> str:
        """Convert cover letter DOCX to PDF."""
        output_path = self.output_dir / f"{output_name}_cover_letter.pdf"

        if DOCX2PDF_AVAILABLE:
            convert(docx_path, str(output_path))
        else:
            try:
                from reportlab.lib.pagesizes import letter
                from reportlab.pdfgen import canvas

                c = canvas.Canvas(str(output_path), pagesize=letter)
                width, height = letter
                y = height - 50
                c.drawString(50, y, f"Cover Letter: {output_name}")
                c.drawString(50, y - 20, "(Please convert DOCX to PDF using external tool)")
                c.save()
            except ImportError:
                raise ImportError(
                    "PDF creation requires either docx2pdf (Windows/Mac) or reportlab. "
                    "Install with: pip install docx2pdf (Windows/Mac) or pip install reportlab"
                )

        return str(output_path)


class ResumeAdaptationSkill:
    """Main skill class that orchestrates the resume adaptation process."""

    def __init__(self):
        self.parser = DocumentParser()
        self.resume_analyzer = ResumeAnalyzer()
        self.job_analyzer = JobDescriptionAnalyzer()
        self.document_builder = None

    def run(self, args) -> Dict[str, Any]:
        """Execute the resume adaptation workflow."""
        results = {
            'success': False,
            'docx_path': None,
            'pdf_path': None,
            'cover_letter_docx': None,
            'cover_letter_pdf': None,
            'ats_score': None,
            'errors': []
        }

        try:
            # Step 1: Parse resume files
            print("📄 Parsing resume(s)...")
            resume_texts = []
            for resume_path in args.resume:
                if not os.path.exists(resume_path):
                    raise FileNotFoundError(f"Resume file not found: {resume_path}")
                text = self.parser.parse(resume_path)
                resume_texts.append(text)

            # Combine multiple resumes (useful if user has different versions)
            combined_resume_text = "\n\n".join(resume_texts)

            # Step 2: Parse job description
            print("💼 Parsing job description...")
            job_text = self._get_job_description(args.job)

            # Step 3: Analyze resume
            print("🔍 Analyzing resume content...")
            parsed_resume = self.resume_analyzer.analyze(combined_resume_text)

            # Step 4: Analyze job requirements
            print("🎯 Extracting job requirements...")
            job_requirements = self.job_analyzer.analyze(job_text)

            # Step 5: ATS Scoring (if requested)
            if getattr(args, 'ats_score', False):
                print("📊 Calculating ATS compatibility score...")
                ats_scorer = ATSScorer(parsed_resume, job_requirements)
                ats_results = ats_scorer.score()
                results['ats_score'] = ats_results
                self._print_ats_report(ats_results)

            # Step 6: Create tailored resume
            print("✨ Tailoring resume for the position...")
            tailor = ResumeTailor(parsed_resume, job_requirements, args.requirements or "")
            tailored_resume = tailor.tailor()

            # Step 7: Build output documents
            print("📝 Creating output documents...")
            self.document_builder = ResumeDocumentBuilder(args.output_dir)

            docx_path = self.document_builder.build_docx(tailored_resume, args.output_name)
            results['docx_path'] = docx_path
            print(f"   ✓ Created: {docx_path}")

            try:
                pdf_path = self.document_builder.build_pdf(docx_path, args.output_name)
                results['pdf_path'] = pdf_path
                print(f"   ✓ Created: {pdf_path}")
            except Exception as e:
                results['errors'].append(f"PDF creation failed: {str(e)}")
                print(f"   ⚠ PDF creation skipped: {str(e)}")

            # Step 8: Generate cover letter (if requested)
            if getattr(args, 'cover_letter', False):
                print("💌 Generating cover letter...")
                cover_gen = CoverLetterGenerator(tailored_resume, job_requirements, args.requirements or "")
                cover_letter_text = cover_gen.generate()

                # Save cover letter as text file for reference
                cover_txt_path = Path(args.output_dir) / f"{args.output_name}_cover_letter.txt"
                with open(cover_txt_path, 'w') as f:
                    f.write(cover_letter_text)
                print(f"   ✓ Created: {cover_txt_path}")

                try:
                    cover_docx_path = self.document_builder.build_cover_letter_docx(cover_letter_text, args.output_name)
                    results['cover_letter_docx'] = cover_docx_path
                    print(f"   ✓ Created: {cover_docx_path}")

                    cover_pdf_path = self.document_builder.build_cover_letter_pdf(cover_docx_path, args.output_name)
                    results['cover_letter_pdf'] = cover_pdf_path
                    print(f"   ✓ Created: {cover_pdf_path}")
                except Exception as e:
                    results['errors'].append(f"Cover letter PDF creation failed: {str(e)}")
                    print(f"   ⚠ Cover letter PDF creation skipped: {str(e)}")

            results['success'] = True
            print("\n✅ Resume adaptation complete!")

        except Exception as e:
            results['errors'].append(str(e))
            print(f"\n❌ Error: {str(e)}")

        return results

    def _print_ats_report(self, ats_results: Dict[str, Any]):
        """Print formatted ATS scoring report."""
        print("\n" + "=" * 50)
        print("📊 ATS COMPATIBILITY REPORT")
        print("=" * 50)
        print(f"\n🎯 Overall Score: {ats_results['overall']}/100")

        print("\n📋 Section Scores:")
        for section, score in ats_results['sections'].items():
            status = "✓" if score >= 70 else "⚠" if score >= 50 else "✗"
            print(f"   {status} {section.replace('_', ' ').title()}: {score}/100")

        print(f"\n📝 Format Score: {ats_results['format_score']}/100")
        print(f"✍️  Content Score: {ats_results['content_score']}/100")

        print("\n💡 Suggestions for Improvement:")
        for i, suggestion in enumerate(ats_results['suggestions'], 1):
            print(f"   {i}. {suggestion}")

        if ats_results['keyword_match'].get('required_missing'):
            print("\n🔍 Missing Required Skills:")
            for skill in ats_results['keyword_match']['required_missing'][:5]:
                print(f"   • {skill}")

        print("=" * 50)

    def _get_job_description(self, job_source: str) -> str:
        """Get job description from text, URL, or file."""
        # Check if it's a URL
        if job_source.startswith(('http://', 'https://')):
            return self._fetch_from_url(job_source)

        # Check if it's a file path
        if os.path.exists(job_source):
            return self.parser.parse(job_source)

        # Assume it's direct text
        return job_source

    def _fetch_from_url(self, url: str) -> str:
        """Fetch job description from a URL."""
        try:
            import requests
            from bs4 import BeautifulSoup

            headers = {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.0'
            }

            response = requests.get(url, headers=headers, timeout=30)
            response.raise_for_status()

            soup = BeautifulSoup(response.text, 'html.parser')

            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()

            # Get text
            text = soup.get_text()

            # Clean up whitespace
            lines = (line.strip() for line in text.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            text = '\n'.join(chunk for chunk in chunks if chunk)

            return text

        except ImportError:
            raise ImportError(
                "URL fetching requires requests and beautifulsoup4. "
                "Install with: pip install requests beautifulsoup4"
            )
        except Exception as e:
            raise Exception(f"Failed to fetch job description from URL: {str(e)}")


def main():
    parser = argparse.ArgumentParser(
        description='Adapt a resume to match a specific job description',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
    # Adapt resume using job description text
    python resume_adapter.py --resume my_resume.pdf --job "Software Engineer position..."

    # Adapt resume from job posting URL
    python resume_adapter.py --resume resume.docx --job https://company.com/jobs/123

    # Adapt with specific requirements
    python resume_adapter.py --resume resume.pdf --job job_desc.txt --requirements "Emphasize leadership experience"

    # Combine multiple resume versions
    python resume_adapter.py --resume version1.pdf version2.docx --job "Job description..."

    # Include ATS scoring
    python resume_adapter.py --resume resume.pdf --job "Job description..." --ats-score

    # Generate cover letter
    python resume_adapter.py --resume resume.pdf --job "Job description..." --cover-letter

    # Full feature set
    python resume_adapter.py --resume resume.pdf --job "Job description..." --ats-score --cover-letter
        """
    )

    parser.add_argument(
        '--resume', '-r',
        nargs='+',
        required=True,
        help='Path(s) to resume file(s) (PDF, DOCX, TXT, MD)'
    )

    parser.add_argument(
        '--job', '-j',
        required=True,
        help='Job description: text string, URL, or file path (PDF, DOCX, TXT)'
    )

    parser.add_argument(
        '--requirements', '-req',
        help='Additional user requirements or preferences'
    )

    parser.add_argument(
        '--output-name', '-n',
        default='tailored_resume',
        help='Base name for output files (default: tailored_resume)'
    )

    parser.add_argument(
        '--output-dir', '-o',
        default='./output',
        help='Directory for output files (default: ./output)'
    )

    parser.add_argument(
        '--ats-score', '-ats',
        action='store_true',
        help='Calculate ATS (Applicant Tracking System) compatibility score'
    )

    parser.add_argument(
        '--cover-letter', '-cl',
        action='store_true',
        help='Generate a tailored cover letter'
    )

    args = parser.parse_args()

    skill = ResumeAdaptationSkill()
    results = skill.run(args)

    sys.exit(0 if results['success'] else 1)


if __name__ == '__main__':
    main()
