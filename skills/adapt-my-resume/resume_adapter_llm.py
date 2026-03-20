#!/usr/bin/env python3
"""
Resume Adaptation Skill - LLM-Native Version

This skill does minimal work (file/URL fetching) and delegates ALL processing
to the underlying LLM (Claude) for superior quality results.

Usage:
    python resume_adapter_llm.py --resume resume.pdf --job "job description" --claude-api-key $ANTHROPIC_API_KEY
    python resume_adapter_llm.py --resume resume.pdf --job https://company.com/job --claude-api-key $ANTHROPIC_API_KEY
"""

import argparse
import os
import re
import sys
from pathlib import Path
from typing import List, Dict, Any, Optional
from dataclasses import dataclass

# Document processing imports - only for extracting text, NOT for analysis
try:
    import docx
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from pypdf import PdfReader
    PDF_AVAILABLE = True
except ImportError:
    try:
        from PyPDF2 import PdfReader
        PDF_AVAILABLE = True
    except ImportError:
        PDF_AVAILABLE = False

try:
    import requests
    from bs4 import BeautifulSoup
    REQUESTS_AVAILABLE = True
except ImportError:
    REQUESTS_AVAILABLE = False

try:
    from anthropic import Anthropic
    ANTHROPIC_AVAILABLE = True
except ImportError:
    ANTHROPIC_AVAILABLE = False


class DocumentFetcher:
    """ONLY fetches/extracts text from documents - no analysis."""

    @staticmethod
    def extract_text_from_pdf(file_path: str) -> str:
        """Extract raw text from PDF."""
        if not PDF_AVAILABLE:
            raise ImportError("PDF parsing requires pypdf. Install with: pip install pypdf")

        reader = PdfReader(file_path)
        text = ""
        for page in reader.pages:
            text += page.extract_text() + "\n"
        return text

    @staticmethod
    def extract_text_from_docx(file_path: str) -> str:
        """Extract raw text from DOCX."""
        if not DOCX_AVAILABLE:
            raise ImportError("DOCX parsing requires python-docx. Install with: pip install python-docx")

        doc = Document(file_path)
        text = ""
        for para in doc.paragraphs:
            text += para.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + "\n"
        return text

    @staticmethod
    def extract_text_from_txt(file_path: str) -> str:
        """Read text from plain text file."""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()

    @classmethod
    def extract_from_file(cls, file_path: str) -> str:
        """Extract text from any supported file type."""
        ext = Path(file_path).suffix.lower()
        extractors = {
            '.pdf': cls.extract_text_from_pdf,
            '.docx': cls.extract_text_from_docx,
            '.txt': cls.extract_text_from_txt,
            '.md': cls.extract_text_from_txt,
        }

        if ext not in extractors:
            raise ValueError(f"Unsupported file format: {ext}")

        return extractors[ext](file_path)

    @staticmethod
    def fetch_url(url: str) -> str:
        """Fetch content from URL."""
        if not REQUESTS_AVAILABLE:
            raise ImportError("URL fetching requires requests and beautifulsoup4. "
                            "Install with: pip install requests beautifulsoup4")

        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
        }

        response = requests.get(url, headers=headers, timeout=30)
        response.raise_for_status()

        soup = BeautifulSoup(response.text, 'html.parser')

        # Remove script and style elements
        for script in soup(["script", "style"]):
            script.decompose()

        # Get text
        text = soup.get_text()

        # Clean up whitespace
        lines = (line.strip() for line in text.splitlines())
        chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
        text = '\n'.join(chunk for chunk in chunks if chunk)

        return text


class ClaudeProcessor:
    """Delegates all processing to Claude via API calls."""

    def __init__(self, api_key: Optional[str] = None):
        if not ANTHROPIC_AVAILABLE:
            raise ImportError("Claude integration requires anthropic SDK. "
                            "Install with: pip install anthropic")

        self.api_key = api_key or os.environ.get('ANTHROPIC_API_KEY')
        if not self.api_key:
            raise ValueError("Anthropic API key required. Set ANTHROPIC_API_KEY env var or pass --claude-api-key")

        self.client = Anthropic(api_key=self.api_key)

    def analyze_and_tailor(self, resume_text: str, job_source: str, user_requirements: str = "") -> Dict[str, Any]:
        """Send everything to Claude and get back all processed results."""

        # Determine if job_source is URL, file, or text
        job_text = job_source
        job_source_type = "text"
        if job_source.startswith(('http://', 'https://')):
            job_source_type = "URL"
        elif os.path.exists(job_source):
            job_source_type = "file"
            job_text = DocumentFetcher.extract_from_file(job_source)

        prompt = f"""You are an expert career coach and professional resume writer. I need you to:

1. ANALYZE the candidate's resume below
2. EXTRACT requirements from the job description
3. TAILOR the resume to match the job
4. GENERATE an ATS compatibility score with analysis
5. WRITE a compelling cover letter

=== CANDIDATE'S RESUME ===
{resume_text}

=== JOB DESCRIPTION ({job_source_type}) ===
{job_text}

=== USER REQUIREMENTS ===
{user_requirements if user_requirements else "None specified - use your best judgment"}

---

Please provide your response in the following format:

## ATS COMPATIBILITY ANALYSIS

**Overall Score:** X/100

**Section Scores:**
- Keyword Match: X/100
- Skills Alignment: X/100
- Experience Relevance: X/100
- Format & Readability: X/100

**Missing Keywords/Skills:**
- List important missing items

**Strengths:**
- What's already well-aligned

**Improvement Suggestions:**
- Specific actionable recommendations

---

## TAILORED RESUME

[Provide the complete tailored resume in professional format with:
- Candidate's name and contact info (preserved from original)
- Professional Summary tailored to the job
- Skills section with relevant skills prioritized
- Experience section highlighting relevant accomplishments
- Education section

Use clear formatting with proper section headers. Make it ATS-friendly but human-readable.]

---

## COVER LETTER

[Provide a compelling cover letter that:
- References the specific job and company
- Opens with a strong hook connecting candidate's background to job
- Highlights 2-3 most relevant achievements/experiences
- Shows enthusiasm and fit
- Closes with a clear call to action

Format as a professional business letter.]

---

## SUMMARY OF CHANGES

[Briefly explain what you changed and why]
"""

        print("🤖 Sending to Claude for analysis and tailoring...")

        response = self.client.messages.create(
            model="claude-3-5-sonnet-20241022",
            max_tokens=8000,
            temperature=0.3,
            messages=[
                {"role": "user", "content": prompt}
            ]
        )

        # Parse the response
        content = response.content[0].text

        return self._parse_claude_response(content)

    def _parse_claude_response(self, content: str) -> Dict[str, Any]:
        """Parse Claude's response into structured sections."""
        result = {
            'ats_analysis': '',
            'tailored_resume': '',
            'cover_letter': '',
            'summary_of_changes': '',
            'raw_response': content
        }

        # Split by section headers
        sections = re.split(r'\n##\s+', content)

        for section in sections:
            section = section.strip()
            if section.startswith('ATS COMPATIBILITY'):
                result['ats_analysis'] = section
            elif section.startswith('TAILORED RESUME'):
                # Remove the header line
                result['tailored_resume'] = '\n'.join(section.split('\n')[1:]).strip()
            elif section.startswith('COVER LETTER'):
                result['cover_letter'] = '\n'.join(section.split('\n')[1:]).strip()
            elif section.startswith('SUMMARY OF CHANGES'):
                result['summary_of_changes'] = '\n'.join(section.split('\n')[1:]).strip()

        return result


class OutputWriter:
    """Writes output files."""

    def __init__(self, output_dir: str):
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)

    def write_outputs(self, results: Dict[str, Any], output_name: str):
        """Write all output files."""
        files_created = []

        # Write tailored resume
        if results.get('tailored_resume'):
            resume_path = self.output_dir / f"{output_name}_resume.txt"
            with open(resume_path, 'w', encoding='utf-8') as f:
                f.write(results['tailored_resume'])
            files_created.append(str(resume_path))
            print(f"   ✓ Created: {resume_path}")

        # Write cover letter
        if results.get('cover_letter'):
            cover_path = self.output_dir / f"{output_name}_cover_letter.txt"
            with open(cover_path, 'w', encoding='utf-8') as f:
                f.write(results['cover_letter'])
            files_created.append(str(cover_path))
            print(f"   ✓ Created: {cover_path}")

        # Write ATS analysis
        if results.get('ats_analysis'):
            ats_path = self.output_dir / f"{output_name}_ats_analysis.txt"
            with open(ats_path, 'w', encoding='utf-8') as f:
                f.write(results['ats_analysis'])
            files_created.append(str(ats_path))
            print(f"   ✓ Created: {ats_path}")

        # Write full response for reference
        if results.get('raw_response'):
            full_path = self.output_dir / f"{output_name}_full_response.txt"
            with open(full_path, 'w', encoding='utf-8') as f:
                f.write(results['raw_response'])
            files_created.append(str(full_path))
            print(f"   ✓ Created: {full_path}")

        return files_created


class ResumeAdaptationSkillLLM:
    """Main skill class - minimal logic, maximum delegation to LLM."""

    def __init__(self, claude_api_key: Optional[str] = None):
        self.fetcher = DocumentFetcher()
        self.processor = ClaudeProcessor(claude_api_key)
        self.writer = None

    def run(self, args) -> Dict[str, Any]:
        """Execute the skill workflow."""
        results = {
            'success': False,
            'files_created': [],
            'errors': []
        }

        try:
            # Step 1: Fetch resume text
            print("📄 Reading resume...")
            resume_texts = []
            for resume_path in args.resume:
                if not os.path.exists(resume_path):
                    raise FileNotFoundError(f"Resume file not found: {resume_path}")
                text = self.fetcher.extract_from_file(resume_path)
                resume_texts.append(text)
                print(f"   ✓ Read: {resume_path}")

            combined_resume = "\n\n".join(resume_texts)

            # Step 2: Fetch job description
            print("💼 Processing job description...")
            job_source = args.job
            if job_source.startswith(('http://', 'https://')):
                print(f"   Fetching from URL: {job_source}")
                job_text = self.fetcher.fetch_url(job_source)
                print(f"   ✓ Fetched {len(job_text)} characters")
            elif os.path.exists(job_source):
                print(f"   Reading from file: {job_source}")
                job_text = self.fetcher.extract_from_file(job_source)
                print(f"   ✓ Read {len(job_text)} characters")
            else:
                print("   Using provided text")
                job_text = job_source
                print(f"   ✓ {len(job_text)} characters")

            # Step 3: DELEGATE EVERYTHING TO CLAUDE
            print("\n🤖 Delegating to Claude for intelligent analysis...")
            processed = self.processor.analyze_and_tailor(
                combined_resume,
                job_source,  # Pass original source so Claude knows if it was URL/file/text
                args.requirements or ""
            )

            # Step 4: Write outputs
            print("\n📝 Writing output files...")
            self.writer = OutputWriter(args.output_dir)
            files = self.writer.write_outputs(processed, args.output_name)
            results['files_created'] = files

            # Print ATS analysis preview
            if processed.get('ats_analysis'):
                print("\n" + "="*60)
                print("📊 ATS ANALYSIS PREVIEW")
                print("="*60)
                # Extract and show overall score
                match = re.search(r'Overall Score:\s*(\d+)/100', processed['ats_analysis'])
                if match:
                    print(f"\n🎯 Overall Score: {match.group(1)}/100")
                print("\n(See full analysis in output files)")

            results['success'] = True
            print("\n✅ Resume adaptation complete!")

        except Exception as e:
            results['errors'].append(str(e))
            print(f"\n❌ Error: {str(e)}")
            import traceback
            traceback.print_exc()

        return results


def main():
    parser = argparse.ArgumentParser(
        description='Resume Adaptation Skill - LLM-Native Version',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
    # Basic usage with text job description
    python resume_adapter_llm.py --resume my_resume.pdf --job "Software Engineer position..." --claude-api-key $KEY

    # Using a job posting URL
    python resume_adapter_llm.py --resume resume.docx --job https://company.com/jobs/123 --claude-api-key $KEY

    # With custom requirements
    python resume_adapter_llm.py --resume resume.pdf --job job_desc.txt --requirements "Emphasize leadership" --claude-api-key $KEY

Environment:
    Set ANTHROPIC_API_KEY environment variable or pass --claude-api-key
        """
    )

    parser.add_argument(
        '--resume', '-r',
        nargs='+',
        required=True,
        help='Path(s) to resume file(s) (PDF, DOCX, TXT, MD)'
    )

    parser.add_argument(
        '--job', '-j',
        required=True,
        help='Job description: text string, URL, or file path'
    )

    parser.add_argument(
        '--requirements', '-req',
        help='Additional user requirements or preferences'
    )

    parser.add_argument(
        '--output-name', '-n',
        default='tailored_resume',
        help='Base name for output files (default: tailored_resume)'
    )

    parser.add_argument(
        '--output-dir', '-o',
        default='./output',
        help='Directory for output files (default: ./output)'
    )

    parser.add_argument(
        '--claude-api-key',
        help='Anthropic API key (or set ANTHROPIC_API_KEY env var)'
    )

    args = parser.parse_args()

    skill = ResumeAdaptationSkillLLM(args.claude_api_key)
    results = skill.run(args)

    sys.exit(0 if results['success'] else 1)


if __name__ == '__main__':
    main()
